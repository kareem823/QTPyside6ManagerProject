from PySide6.QtWidgets import QMessageBox, QFileDialog
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

class InvoiceExporter:
    def __init__(self, parent=None):
        self.parent = parent  # Reference to the InvoicePage if needed

    def export_invoice_pdf(self, invoice_data):
        """
        Exports the invoice data to a PDF file.
        """
        try:
            # Open a file dialog to select the save location
            file_path, _ = QFileDialog.getSaveFileName(
                self.parent,
                "Save Invoice as PDF",
                "",
                "PDF Files (*.pdf);;All Files (*)"
            )

            if not file_path:
                QMessageBox.warning(self.parent, "Warning", "Export cancelled.")
                return

            # Create a PDF document
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            elements = []

            # Define styles
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(name='CenterBold', alignment=1, fontSize=16, leading=20, spaceAfter=20, fontName='Helvetica-Bold'))
            styles.add(ParagraphStyle(name='LeftNormal', alignment=0, fontSize=12, leading=14, spaceAfter=12))
            styles.add(ParagraphStyle(name='LeftItalic', alignment=0, fontSize=12, leading=14, spaceAfter=12, fontName='Helvetica-Oblique'))

            # Business Name
            elements.append(Paragraph(invoice_data['business_name'], styles['CenterBold']))

            # Invoice Title
            elements.append(Paragraph("INVOICE", styles['CenterBold']))
            elements.append(Spacer(1, 12))

            # Invoice Details
            details = [
                ["Invoice Number:", invoice_data['invoice_number']],
                ["Customer Name:", invoice_data['customer_name']],
                ["Customer Phone", invoice_data['customer_phone']],
                ["Customer Email:", invoice_data['customer_email']],
                ["Customer Address:", invoice_data['customer_address']],
                ["Invoice Date:", invoice_data['invoice_date']],
                ["Due Date:", invoice_data['due_invoice_date']]
            ]
            table = Table(details, colWidths=[120, 300])
            table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 12),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            elements.append(table)
            elements.append(Spacer(1, 12))

            # Invoice Items Header
            items_header = ["Description", "Quantity", "Price", "Total"]
            items = [items_header]

            # Invoice Items Data
            for item in invoice_data['invoice_items']:
                total_price = item['quantity'] * item['price']
                items.append([
                    item['item_description'],
                    str(item['quantity']),
                    f"${item['price']:.2f}",
                    f"${total_price:.2f}"
                ])

            # Define table style for items
            table = Table(items, colWidths=[250, 70, 70, 70])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.gray),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('FONTSIZE', (0, 1), (-1, -1), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ]))

            # Optional: Add alternating row colors for better readability
            for i in range(1, len(items)):
                if i % 2 == 0:
                    bg_color = colors.whitesmoke
                else:
                    bg_color = colors.lightgrey
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, i), (-1, i), bg_color)
                ]))

            elements.append(table)
            elements.append(Spacer(1, 12))

            # Total Amount Due
            total_due = sum(item['quantity'] * item['price'] for item in invoice_data['invoice_items'])
            total_paragraph = Paragraph(f"<b>Total Amount Due: ${total_due:.2f}</b>", styles['LeftNormal'])
            elements.append(total_paragraph)

            # Notes
            if invoice_data['notes']:
                elements.append(Paragraph("Notes:", styles['LeftItalic']))
                elements.append(Paragraph(invoice_data['notes'], styles['LeftItalic']))
                elements.append(Spacer(1, 11))

            # Company Footer
            if invoice_data['company_footer']:
                #I wanna add a line to separate the footer
                elements.append(Spacer(1, 11))
                elements.append(Paragraph(invoice_data['company_footer'], styles['CenterBold']))
                elements.append(Spacer(1, 8))


            # Build the PDF
            doc.build(elements)

            QMessageBox.information(self.parent, "Success", "PDF exported successfully.")
        except Exception as e:
            QMessageBox.critical(self.parent, "Error", f"Failed to export PDF.\nError: {str(e)}")

    def export_invoice_word(self, invoice_data):
        """
        Exports the invoice data to a Word document.
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # Open a file dialog to select the save location
            file_path, _ = QFileDialog.getSaveFileName(
                self.parent,
                "Save Invoice as Word Document",
                "",
                "Word Documents (*.docx);;All Files (*)"
            )

            if not file_path:
                QMessageBox.warning(self.parent, "Warning", "Export cancelled.")
                return

            document = Document()

            # Add Business Name
            heading = document.add_heading(invoice_data['business_name'], 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Invoice Title
            invoice_title = document.add_heading('INVOICE', level=1)
            invoice_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Invoice Details
            table = document.add_table(rows=4, cols=2)
            table.style = 'Light List Accent 1'
            table.autofit = True

            details = [
                ["Invoice Number:", invoice_data['invoice_number']],
                ["Customer Name:", invoice_data['customer_name']],
                ["Customer Phone:", invoice_data['customer_phone']],
                ["Customer Email:", invoice_data['customer_email']],
                ["Customer Address:", invoice_data['customer_address']],
                ["Invoice Date:", invoice_data['invoice_date']],
                ["Due Date:", invoice_data['due_invoice_date']]
            ]

            for row, detail in zip(table.rows, details):
                row.cells[0].text = detail[0]
                row.cells[1].text = detail[1]

            document.add_paragraph()

            # Invoice Items
            items_table = document.add_table(rows=1, cols=4)
            items_table.style = 'Light List Accent 1'
            hdr_cells = items_table.rows[0].cells
            hdr_cells[0].text = 'Description'
            hdr_cells[1].text = 'Quantity'
            hdr_cells[2].text = 'Price'
            hdr_cells[3].text = 'Total'

            for item in invoice_data['invoice_items']:
                row_cells = items_table.add_row().cells
                row_cells[0].text = item['item_description']
                row_cells[1].text = str(item['quantity'])
                row_cells[2].text = f"${item['price']:.2f}"
                total_price = item['quantity'] * item['price']
                row_cells[3].text = f"${total_price:.2f}"

            document.add_paragraph()

                        # Total Amount Due
            total_due = sum(item['quantity'] * item['price'] for item in invoice_data['invoice_items'])
            total_paragraph = document.add_paragraph()
            run = total_paragraph.add_run(f"Total Amount Due: ${total_due:.2f}")
            run.bold = True

            # Notes
            if invoice_data['notes']:
                document.add_heading('Additional Notes:', level=2)
                document.add_paragraph(invoice_data['notes'])
            
            # Company Footer
            if invoice_data['company_footer']:
                document.add_paragraph(invoice_data['company_footer'])

            # Save the document
            document.save(file_path)

            QMessageBox.information(self.parent, "Success", "Word document exported successfully.")
        except Exception as e:
            QMessageBox.critical(self.parent, "Error", f"Failed to export Word document.\nError: {str(e)}")
